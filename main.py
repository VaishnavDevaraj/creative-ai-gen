import os
import io
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from groq import Groq
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from fpdf import FPDF

load_dotenv()
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class GenerationRequest(BaseModel):
    prompt: str
    style: str

class DownloadRequest(BaseModel):
    text: str

def create_formatted_docx(text: str) -> io.BytesIO:
    document = Document()
    lines = [line for line in text.split('\n') if line.strip()]
    if not lines:
        document.add_paragraph(text)
    else:
        document.add_heading(lines[0], level=1)
        for line in lines[1:]:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.space_after = Inches(0.15)
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- THIS FUNCTION IS UPDATED FOR DEPLOYMENT ---
def create_formatted_pdf(text: str) -> io.BytesIO:
    """Creates a formatted PDF document in memory using bundled Unicode fonts."""
    pdf = FPDF()
    
    # Use the font files we bundled with the application
    font_path_regular = os.path.join(os.path.dirname(__file__), 'fonts', 'arial.ttf')
    font_path_bold = os.path.join(os.path.dirname(__file__), 'fonts', 'arialbd.ttf')
    
    pdf.add_font('Arial', '', font_path_regular)
    pdf.add_font('Arial', 'B', font_path_bold)
    
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    lines = [line for line in text.split('\n') if line.strip()]

    if not lines:
        pdf.set_font('Arial', size=12)
        pdf.multi_cell(0, 10, text)
    else:
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, lines[0], ln=True, align='C')
        pdf.ln(10)
        pdf.set_font('Arial', size=12)
        for line in lines[1:]:
            pdf.multi_cell(0, 10, line)
            pdf.ln(5)
    
    pdf_bytes = pdf.output(dest='S')
    return io.BytesIO(pdf_bytes)

# --- Endpoints are unchanged ---
@app.get("/")
def read_root():
    return {"status": "Creative AI Backend is running!"}

@app.post("/api/generate/text")
async def generate_text(request: GenerationRequest):
    # ... same code as before
    try:
        groq_client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    except Exception as e:
        return {"error": f"Failed to initialize Groq client: {e}"}, 500
    system_prompt = f"You are an expert creative assistant. Your task is to generate content based on the user's request. The desired style is: {request.style}. Format your response with clear paragraph breaks."
    try:
        chat_completion = groq_client.chat.completions.create(messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": request.prompt}], model="llama3-8b-8192", temperature=0.7, max_tokens=1024)
        response_content = chat_completion.choices[0].message.content
        return {"data": response_content}
    except Exception as e:
        return {"error": f"An error occurred with the AI service: {e}"}, 503

@app.post("/api/download/docx")
async def download_docx(request: DownloadRequest):
    file_stream = create_formatted_docx(request.text)
    headers = {'Content-Disposition': 'attachment; filename="creative_document.docx"'}
    return StreamingResponse(file_stream, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)

@app.post("/api/download/pdf")
async def download_pdf(request: DownloadRequest):
    file_stream = create_formatted_pdf(request.text)
    headers = {'Content-Disposition': 'attachment; filename="creative_document.pdf"'}
    return StreamingResponse(file_stream, media_type='application/pdf', headers=headers)